"""DOCX export for tailored CVs."""
from __future__ import annotations

import io
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

from auth.config import current_active_verified_user
from db.base import get_session
from db.models import MasterCV, TailoredCV, User

router = APIRouter()


def _build_docx(content_json: str) -> bytes:
    data = json.loads(content_json)
    doc = Document()

    # Compact page margins
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # Name header
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(data.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if data.get("title"):
        title_para = doc.add_paragraph(data["title"])
        title_para.runs[0].font.size = Pt(11)
        title_para.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Contact line
    contact = data.get("contact", {})
    contact_parts = [
        contact.get("email"), contact.get("phone"),
        contact.get("portfolio"), contact.get("github"), contact.get("location"),
    ]
    contact_str = "  |  ".join(p for p in contact_parts if p)
    if contact_str:
        cp = doc.add_paragraph(contact_str)
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    for section in data.get("sections", []):
        heading = section.get("heading", "")
        stype = section.get("type", "paragraph")

        h = doc.add_paragraph(heading.upper())
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(9)
        h.runs[0].font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        h_format = h.paragraph_format
        h_format.space_before = Pt(6)
        h_format.space_after = Pt(2)
        h.runs[0].font.all_caps = True

        if stype == "paragraph" and section.get("content"):
            raw = section["content"].replace("[AI:]", "").replace("[:AI]", "")
            p = doc.add_paragraph(_strip_bold_markers(raw))
            p.runs[0].font.size = Pt(10)

        elif stype == "bullets":
            for item in section.get("items", []):
                raw = item.replace("[AI:]", "").replace("[:AI]", "")
                p = doc.add_paragraph(style="List Bullet")
                _add_bold_run(p, raw)

        elif stype == "entries":
            for entry in section.get("entries", []):
                org_para = doc.add_paragraph()
                org_run = org_para.add_run(entry.get("org", ""))
                org_run.bold = True
                org_run.font.size = Pt(10)
                org_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
                if entry.get("date"):
                    date_run = org_para.add_run(f"  {entry['date']}")
                    date_run.font.size = Pt(9)
                    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                if entry.get("role"):
                    rp = doc.add_paragraph(entry["role"])
                    rp.runs[0].italic = True
                    rp.runs[0].font.size = Pt(10)
                for bullet in entry.get("bullets", []):
                    raw = bullet.replace("[AI:]", "").replace("[:AI]", "")
                    bp = doc.add_paragraph(style="List Bullet")
                    _add_bold_run(bp, raw)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _strip_bold_markers(text: str) -> str:
    import re
    return re.sub(r"\*\*([^*]+)\*\*", r"\1", text)


def _add_bold_run(para, text: str):
    import re
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)
    for run in para.runs:
        run.font.size = Pt(10)


@router.get("/{master_cv_id}/tailored/{tailored_cv_id}/docx")
async def download_docx(
    master_cv_id: int,
    tailored_cv_id: int,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_active_verified_user),
):
    cv = await session.get(MasterCV, master_cv_id)
    if not cv or cv.user_id != user.id:
        raise HTTPException(404, "CV not found")

    tailored = await session.get(TailoredCV, tailored_cv_id)
    if not tailored or tailored.master_cv_id != master_cv_id:
        raise HTTPException(404, "Tailored CV not found")

    docx_bytes = _build_docx(tailored.content_json)
    filename = f"tailored-cv-{tailored_cv_id}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
